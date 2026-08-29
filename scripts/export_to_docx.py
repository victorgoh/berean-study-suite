#!/usr/bin/env python3
"""
Convert markdown study files into professionally styled Microsoft Word (.docx) documents.
"""

import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def markdown_to_docx(md_text: str, output_path: str, title: str = "Berean Study Manuscript"):
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Style Settings
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "F4F5F7")
                set_cell_margins(cell, 120, 120, 150, 150)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                run = p.add_run('\n'.join(code_block_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
                doc.add_paragraph() # spacing
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Markdown Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
                # skip separator row (table_lines[1])
                data_rows = []
                for row_line in table_lines[2:]:
                    row_data = [c.strip() for c in row_line.strip('|').split('|')]
                    data_rows.append(row_data)
                
                table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Format Header
                for col_idx, h_text in enumerate(headers):
                    cell = table.cell(0, col_idx)
                    set_cell_background(cell, "1E3A8A") # Navy Blue
                    set_cell_margins(cell, 120, 120, 150, 150)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run(h_text.replace('**', ''))
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
                
                # Format Data Rows
                for row_idx, row_data in enumerate(data_rows):
                    bg_color = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
                    for col_idx in range(len(headers)):
                        cell = table.cell(row_idx + 1, col_idx)
                        set_cell_background(cell, bg_color)
                        set_cell_margins(cell, 100, 100, 120, 120)
                        val = row_data[col_idx] if col_idx < len(row_data) else ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        # Render inline bold/italic
                        add_formatted_runs(p, val, base_size=Pt(9.5))
                
                doc.add_paragraph() # Spacing
            continue

        # Headings
        if stripped.startswith('# '):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(stripped[2:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Blue
            i += 1
            continue

        if stripped.startswith('## '):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[3:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF) # Blue Accent
            i += 1
            continue

        if stripped.startswith('### '):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(3)
            run = h.add_run(stripped[4:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51) # Slate Grey
            i += 1
            continue

        # Blockquote / Scripture Quotes
        if stripped.startswith('> '):
            quote_text = stripped[2:].strip()
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            set_cell_background(cell, "F0FDF4") # Pale Green / Pale Slate
            set_cell_margins(cell, 120, 120, 160, 160)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, quote_text, base_size=Pt(10.5), italic=True)
            doc.add_paragraph() # Spacing
            i += 1
            continue

        # Bullet points
        if stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('• '):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, bullet_text)
            i += 1
            continue

        # Numbered list
        match_num = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if match_num:
            item_text = match_num.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, item_text)
            i += 1
            continue

        # Horizontal Rule
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E5E7EB"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)
            i += 1
            continue

        # Standard Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_runs(p, stripped)
        i += 1

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated docx: {output_path}")

def add_formatted_runs(paragraph, text, base_size=Pt(11), italic=False):
    # Regex split by bold (**...**) and italic (*...*)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = base_size
            if italic:
                run.italic = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = base_size
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = base_size
            if italic:
                run.italic = True

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python export_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        md = f.read()
    
    markdown_to_docx(md, sys.argv[2])
